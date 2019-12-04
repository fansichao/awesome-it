u""" 细胞溶填写模板 

环境依赖:
    - Python3
    - pip install python-docx xlrd
参考链接:
    - https://python-docx.readthedocs.io/en/latest/index.html


操作顺序：
    终端-> 新建终端
运行命令:
    python K_word.py
注意事项：
    需要关闭 生成的word文档，否则会报错 IOError: [Errno 13] Permission denied: u'a.docx'
"""
import datetime
import os
import sys
import traceback

import docx
import xlrd
from xlrd import xldate_as_tuple

# reload(sys)
# sys.setdefaultencoding('utf-8')


def read_xlsx(file_name, sheet_name=None):
    """ 读取xlsx文件内容
    @param file_name: 文件名称
    @param sheet_name: 表格名称，None时，默认取第一个
    """
    # try:
    #     file_name = file_name.decode('utf-8')
    #     sheet_name = sheet_name.decode('utf-8')
    # except Exception:
    #     print(traceback.print_exc())

    # rbook = xlrd.open_workbook(file_name,formatting_info=True) # 保留excel样式
    rbook = xlrd.open_workbook(file_name)
    # 所有 sheet_name
    sheet_names = rbook.sheet_names()
    if not bool(sheet_name):
        sheet_name = sheet_names[0]
    if sheet_name not in sheet_names:
        print(">>>> 表名称[%s]不存在文件[%s]所有表[%s]中" %
              (sheet_name, file_name, sheet_names))
        return

    sheet = rbook.sheet_by_name(sheet_name)
    rows = sheet.nrows
    cols = sheet.ncols
    all_content = []
    for i in range(rows):
        row_content = []
        for j in range(cols):
            ctype = sheet.cell(i, j).ctype  # 表格的数据类型
            cell = sheet.cell_value(i, j)
            if ctype == 2 and cell % 1 == 0:  # 如果是整形
                cell = int(cell)
            elif ctype == 3:
                # 转成datetime对象
                date = datetime.datetime(*xldate_as_tuple(cell, 0))
                a = date.strftime('%Y-%m-%d %H:%M:%S')
                cell = str(a).split(' ')[0]
                # cell = date.strftime('%Y/%d/%m %H:%M:%S')
            elif ctype == 4:
                cell = True if cell == 1 else False
            row_content.append(cell)
        all_content.append(row_content)
        # print '[' + ','.join("'" + str(element) + "'" for element in row_content) + ']'
    return all_content


def read_docx(file_name):
    doc = docx.Document(file_name)
    #content = '\n'.join([para.text for para in doc.paragraphs])
    # print content
    return doc


def save_docx(doc, out_filename):
    doc.save(out_filename)


filenames = []


def deal_docx(excel_model_name, docx_model_name):
    u" doc处理"
    # xls_data 数据格式 [[],[]]
    xls_data = read_xlsx(excel_model_name, sheet_name='data')
    xls_header = xls_data[0]
    child_data = read_xlsx(excel_model_name, sheet_name='data_sortno')
    child_header = child_data[0]

    def mid_tmp(xls_data, header):
        rows = []
        for index in range(len(xls_data)):
            dic = {}
            if index == 0 or index == 1:
                continue
            for i in range(len(header)):
                dic[header[i]] = xls_data[index][i]
            rows.append(dic)
        return rows

    # xls_dic 数据格式 [{},{}]
    xls_dic = mid_tmp(xls_data, xls_header)
    child_dic = mid_tmp(child_data, child_header)
    child_model = [
        # ①②③④⑤⑥⑦⑧⑨⑩ ⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳
        '①  M1ml+M2ml(M3%DMSO)细胞数:M4×108/ml 细胞总数:M5×108   M6×108/kg体重',
        '②  M1ml+M2ml(M3%DMSO)细胞数:M4×108/ml 细胞总数:M5×108   M6×108/kg体重',
        '③  M1ml+M2ml(M3%DMSO)细胞数:M4×108/ml 细胞总数:M5×108   M6×108/kg体重',
        '④  M1ml+M2ml(M3%DMSO)细胞数:M4×108/ml 细胞总数:M5×108   M6×108/kg体重',
        '⑤  M1ml+M2ml(M3%DMSO)细胞数:M4×108/ml 细胞总数:M5×108   M6×108/kg体重',
        '⑥  M1ml+M2ml(M3%DMSO)细胞数:M4×108/ml 细胞总数:M5×108   M6×108/kg体重',
        '⑦  M1ml+M2ml(M3%DMSO)细胞数:M4×108/ml 细胞总数:M5×108   M6×108/kg体重',
        '⑧  M1ml+M2ml(M3%DMSO)细胞数:M4×108/ml 细胞总数:M5×108   M6×108/kg体重',
        '⑨  M1ml+M2ml(M3%DMSO)细胞数:M4×108/ml 细胞总数:M5×108   M6×108/kg体重',
        '⑩  M1ml+M2ml(M3%DMSO)细胞数:M4×108/ml 细胞总数:M5×108   M6×108/kg体重',
    ]

    import copy
    # 主表数据循环
    for data in xls_dic:
        doc = read_docx(docx_model_name)
        # 获取 子表 数据
        tmp_row = []
        for r in child_dic:
            if r['data_sortno'] == data['sortno']:
                tmp_row.append(r)

        num_len = len(tmp_row)
        child_model_data = copy.deepcopy(child_model)[0:num_len]
        # 子表数据处理
        for i in range(len(tmp_row)):
            for k in tmp_row[i]:
                if k not in ['M1', 'M2', 'M3', 'M4', 'M5', 'M6']:
                    continue
                child_model_data[i] = child_model_data[i].replace(
                    k, str(tmp_row[i][k]))

        while len(child_model_data) <= 10:
            child_model_data.append('')

        for i in range(1, len(child_model_data) + 1):
            data.update({'KK' + str(i): child_model_data[i - 1]})

        # 数据处理
        if data['iscreate'] == 'no':
            continue

       # 表格处理
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for k in data.keys():
                        cell.text = cell.text.replace(k, str(data[k]))

        # 段落处理
        for para in doc.paragraphs:
            for k in data.keys():
                if k in para.text:

                    if 'KK' in para.text:
                        # TODO 样式处理
                        # print para.style
                        pass

                    para.text = para.text.replace(k, str(data[k]))

                    # 单独 指定子表格式

                    #para.style = 'List Bullet'
                    #doc.add_paragraph('Lorem ipsum dolor sit amet.', style='ListBullet')

        save_docx(doc, data['filename'])
        filenames.append(data['filename'])


def main(excel_model_name, docx_model_name):
    deal_docx(excel_model_name, docx_model_name)


if __name__ == '__main__':
    print(u">>>>  请确保 MyModel.xlsx model.docx 是在此目录！！！")
    excel_model_name = "MyModel.xlsx"
    docx_model_name = 'model.docx'
    print(u">> 模板填充 开始,输入文件[%s],[%s]" % (excel_model_name, docx_model_name))

    a = datetime.datetime.now()
    main(excel_model_name, docx_model_name)
    b = datetime.datetime.now()
    print(u">> 模板填充 完成，生成文件[%s],耗时[%s]" % (str(filenames), str(b - a)))
