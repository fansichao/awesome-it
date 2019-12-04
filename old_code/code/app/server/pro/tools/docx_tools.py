#! -*- coding:utf-8 -*-
u""" Word doc/docx 文档处理工具

简单 docx文档使用

*   python-docx：读取，查询以及修改 Microsoft Word 2007/2008 docx 文件。[官网](https://github.com/python-openxml/python-docx)



"""
import os
import logging

from docx import Document
from docx.shared import Inches

import traceback



def get_docx_info(document):
    u""" 获取文档信息
    """
    pass

def read_docx(file_path=None):
    u""" 读取文档

    文档不存在时,创建空白文档
    """
    try:
        document = Document(file_path)
        logging.info(u">> 文档[%s]读取成功",file_path)
    except Exception as e:
        logging.info(e)
        document = Document()
        logging.info(u">> 文档[%s]读取失败,新建空白文档",file_path)
    return document 


def edit_docx(document):
    u""" 编辑文档文本内容
    """
    # 添加段落
    p = document.add_paragraph('A plain paragraph having some ')
    document.add_paragraph('Intense quote', style='Intense Quote')
    p.insert_paragraph_before('Lorem ipsum')
    # 添加标题
    document.add_heading('The REAL meaning of the universe')
    document.add_heading('The role of dolphins', level=2)
    # 添加分页
    document.add_page_break()
    # 添加表
    table = document.add_table(rows=2, cols=2)
    cell = table.cell(0, 1)
    cell.text = 'parrot, possibly dead'
    
    # 指定表格样式
    table.style = 'LightShading-Accent1'
    
    
    # 指定图片尺寸
    pic_path = False
    if pic_path and os.path.exists(pic_path):
        document.add_picture(pic_path, width=Inches(1.0))
    
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )
    
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    
    # 文档保存
    document.save('demo.docx')


if __name__ == '__main__':
    pass
    document = read_docx()
    edit_docx(document)

