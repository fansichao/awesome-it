'''
@Author: Scfan
@Date: 2019-01-11 16:17:44
@LastEditors: Scfan
@LastEditTime: 2019-01-11 16:36:37
@Description: 工作&amp;学习&amp;生活
@Email: 643566992@qq.com
@Company: 上海
@version: V1.0
'''
#! -*- coding:utf-8 -*-
u"""

    Word Api

pip install python-docx
https://python-docx.readthedocs.io/en/latest/index.html
"""
import datetime
import os
import sys
import traceback

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

reload(sys)
sys.setdefaultencoding('utf-8')


def read_docx(file_name):
    '''
    @description: 读取word文档
    @param file_name: word文件名称
    @return: 
    '''
    doc = docx.Document(file_name)
    content = '\n'.join([para.text for para in doc.paragraphs])
    print content
    return doc

def save_docx(doc, out_filename='out.docx'):
    '''
    @description: 保存word文档
    @param doc: word对象
    @param out_filename:输出文件名称
    @return: 
    '''
    doc.save(out_filename)

def deal_docx(doc, rows):
    '''
    @description: Word处理
    @param {type} 
    @return: 
    '''
    for data in rows:
        if data['iscreate'] == 'no':
            continue
        for para in doc.paragraphs:
            for k in data.keys():
                if k in para.text:
                    para.text = para.text.replace(k, str(data[k]))

        save_docx(doc, data['filename'])

def main():
    '''
    @description: 函数入口
    @param {type} 
    @return: 
    '''
    file_name = 'model.docx'
    doc = read_docx(file_name)
    deal_docx(doc, rows)


def docx_demo():
    '''
    @description: 官方demo代码
        链接：https://python-docx.readthedocs.io/en/latest/index.html
    @param {type} 
    @return: 
    '''
    from docx import Document
    from docx.shared import Inches
    # 初始化文档
    document = Document()
    # 添加标题
    document.add_heading('Document Title', 0)
    # 添加段落
    p = document.add_paragraph('A plain paragraph having some ')
    # 在段落中添加数据  粗体、斜体
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    # 应用字符样式
    p.add_run('text with emphasis1.', 'Emphasis')
    run = p.add_run('text with emphasis2.')
    run.style = 'Emphasis'
    p.add_run('italic.').italic = True
    # 设置标题级别
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )
    # 添加图片
    document.add_picture('monty-truth.png', width=Inches(1.25))
    # 添加表格数据
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    # 添加分页符
    document.add_page_break()
    # 文档保存
    document.save('demo.docx')


if __name__ == '__main__':
    
    docx_demo()
    exit()
    main()

